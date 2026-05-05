"""Write a Script or Analysis out as .txt and/or .docx."""
from __future__ import annotations

from pathlib import Path

from .extractor import Script


def _format_timestamp(seconds: float) -> str:
    seconds = int(seconds)
    h, rem = divmod(seconds, 3600)
    m, s = divmod(rem, 60)
    return f"{h:02d}:{m:02d}:{s:02d}" if h else f"{m:02d}:{s:02d}"


def _header(script: Script) -> str:
    kind = "auto-generated" if script.is_generated else "uploader-provided"
    if script.is_translated:
        kind += ", translated by YouTube"
    return (
        f"YouTube video: https://www.youtube.com/watch?v={script.video_id}\n"
        f"Language: {script.language_name} ({script.language})\n"
        f"Source: {kind}\n"
        f"{'-' * 60}\n"
    )


def write_txt(script: Script, path: Path, with_timestamps: bool = True) -> Path:
    lines = [_header(script)]
    if with_timestamps:
        for seg in script.segments:
            ts = _format_timestamp(seg.get("start", 0))
            text = seg.get("text", "").strip()
            if text:
                lines.append(f"[{ts}] {text}")
    else:
        lines.append(script.text)
    path.write_text("\n".join(lines), encoding="utf-8")
    return path


def write_docx(script: Script, path: Path, with_timestamps: bool = True) -> Path:
    try:
        from docx import Document
    except ImportError as e:
        raise RuntimeError(
            "python-docx is not installed. Run: pip install python-docx"
        ) from e

    doc = Document()
    doc.add_heading(f"Transcript - {script.video_id}", level=1)
    doc.add_paragraph(f"URL: https://www.youtube.com/watch?v={script.video_id}")
    doc.add_paragraph(f"Language: {script.language_name} ({script.language})")
    kind = "auto-generated" if script.is_generated else "uploader-provided"
    if script.is_translated:
        kind += ", translated by YouTube"
    doc.add_paragraph(f"Source: {kind}")
    doc.add_paragraph("")

    if with_timestamps:
        for seg in script.segments:
            ts = _format_timestamp(seg.get("start", 0))
            text = seg.get("text", "").strip()
            if text:
                doc.add_paragraph(f"[{ts}] {text}")
    else:
        doc.add_paragraph(script.text)

    doc.save(path)
    return path


def write_analysis_txt(analysis, path: Path) -> Path:
    sep = "=" * 60
    lines = [
        f"YouTube video: https://www.youtube.com/watch?v={analysis.video_id}",
        f"Language: {analysis.language}",
        f"Word count: {analysis.word_count:,}  |  Estimated watch time: {analysis.estimated_watch_minutes} min",
        f"Analysis method: {analysis.method}",
        sep, "",
        "SUMMARY", "-" * 40, "",
        analysis.summary, "",
        sep, "",
        "KEY POINTS", "-" * 40, "",
    ]
    for point in analysis.key_points:
        lines.append(f"  • {point}")
    lines += [
        "", sep, "",
        "DEEP DIVE", "-" * 40, "",
        analysis.deep_dive, "",
        sep, "",
        "AI COMMENTS", "-" * 40, "",
        analysis.ai_comments, "",
        sep, "",
        "TOPICS", "-" * 40, "",
        "  " + ", ".join(analysis.topics),
    ]
    path.write_text("\n".join(lines), encoding="utf-8")
    return path


def write_analysis_docx(analysis, path: Path) -> Path:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    title = doc.add_heading(f"Video Analysis Report", level=1)
    doc.add_paragraph(f"URL: https://www.youtube.com/watch?v={analysis.video_id}")
    doc.add_paragraph(f"Language: {analysis.language}")
    doc.add_paragraph(
        f"Word count: {analysis.word_count:,}  |  "
        f"Estimated watch time: {analysis.estimated_watch_minutes} min"
    )
    doc.add_paragraph(f"Analysis method: {analysis.method}")
    doc.add_paragraph("")

    doc.add_heading("Summary", level=2)
    for para in analysis.summary.split("\n\n"):
        if para.strip():
            doc.add_paragraph(para.strip())
    doc.add_paragraph("")

    doc.add_heading("Key Points", level=2)
    for point in analysis.key_points:
        doc.add_paragraph(point, style="List Bullet")
    doc.add_paragraph("")

    doc.add_heading("Deep Dive", level=2)
    for para in analysis.deep_dive.split("\n\n"):
        if para.strip():
            doc.add_paragraph(para.strip())
    doc.add_paragraph("")

    doc.add_heading("AI Comments", level=2)
    for para in analysis.ai_comments.split("\n\n"):
        if para.strip():
            doc.add_paragraph(para.strip())
    doc.add_paragraph("")

    doc.add_heading("Topics", level=2)
    doc.add_paragraph(", ".join(analysis.topics))

    doc.save(path)
    return path


def output_paths(out_dir: Path, script: Script, suffix: str) -> dict[str, Path]:
    base = f"{suffix}_{script.language}"
    return {"txt": out_dir / f"{base}.txt", "docx": out_dir / f"{base}.docx"}
